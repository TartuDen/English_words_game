import docx


class ReadDocx():
    def __init__(self) -> None:
        pass
        # Open the Word file
        doc = docx.Document('wordFiles/MainEnglish.docx')
        # doc = docx.Document('.\\wordFiles\\MainEnglish voicetest.docx')
        # Get the first table in the document
        self.table = doc.tables[0]
        # self.read_words()
    def read_words(self):
        final_dic_eng_rus = {}
        final_dic_eng_mean = {}
        # Get the second column of the table
        column2 = self.table.columns[1]
        column3 = self.table.columns[2]
        column4 = self.table.columns[3]
        list_ = ["–","-"]
        # Loop through each cell in the column
        count=0
        for cell_eng,cell_rus, cell_mean in zip(column2.cells,column3.cells,column4.cells):
            # Get the text from the cell
            text_eng = cell_eng.text.strip()
            text_rus = cell_rus.text.rstrip("-")
            meaning_eng = cell_mean.text.strip()
            # Find the index of the "-/" characters in the text
            for i in list_:
                index = text_eng.find(i)
                if index > -1:
                    break
            # If the characters are found, print out the text before them
            if index != -1:
                final_dic_eng_rus[text_eng[:index].strip()]=text_rus
                final_dic_eng_mean[text_eng[:index].strip()] = meaning_eng
                count+=1
        return(final_dic_eng_rus,final_dic_eng_mean)

# s=ReadDocx()